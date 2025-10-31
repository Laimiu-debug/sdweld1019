"""
文档导出服务
处理WPS/PQR/pPQR文档导出为Word和PDF
"""
from sqlalchemy.orm import Session
from typing import Optional, BinaryIO
import io
from datetime import datetime
from bs4 import BeautifulSoup

# 注意：这些库需要安装
# pip install python-docx weasyprint beautifulsoup4
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，Word导出功能不可用")

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    print(f"警告: weasyprint不可用，PDF导出功能不可用 ({str(e)[:100]})")

from app.models.wps import WPS
from app.models.pqr import PQR
from app.models.ppqr import PPQR


class DocumentExportService:
    """文档导出服务类"""
    
    def __init__(self, db: Session):
        self.db = db
    
    def export_wps_to_word(self, wps: WPS, style: str = "blue_white") -> io.BytesIO:
        """
        导出WPS为Word文档

        Args:
            wps: WPS对象
            style: 表格风格，可选值：
                - "blue_white": 蓝白相间风格（默认）
                - "plain": 纯白风格
                - "classic": 经典风格（深蓝标题）

        Returns:
            Word文档的BytesIO流
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx未安装，请运行: pip install python-docx")

        # 创建Word文档
        doc = Document()

        # 设置页面（A4）
        section = doc.sections[0]
        section.page_height = Inches(11.69)  # A4高度
        section.page_width = Inches(8.27)    # A4宽度
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)

        # 获取HTML内容
        html_content = wps.document_html or self._generate_default_html(wps)

        # 解析HTML并转换为Word
        self._html_to_word(doc, html_content, style=style)
        
        # 添加页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"打印日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
    
    def export_wps_to_pdf(self, wps: WPS) -> io.BytesIO:
        """
        导出WPS为PDF文档
        
        Args:
            wps: WPS对象
            
        Returns:
            PDF文档的BytesIO流
        """
        if not WEASYPRINT_AVAILABLE:
            raise ImportError("weasyprint未安装，请运行: pip install weasyprint")
        
        # 获取HTML内容
        html_content = wps.document_html or self._generate_default_html(wps)
        
        # 生成完整的HTML文档
        full_html = self._generate_pdf_html(wps, html_content)
        
        # 生成PDF
        pdf_bytes = HTML(string=full_html).write_pdf()
        
        return io.BytesIO(pdf_bytes)
    
    def _html_to_word(self, doc: Document, html_content: str, style: str = "blue_white"):
        """
        将HTML内容转换为Word文档

        Args:
            doc: Word文档对象
            html_content: HTML内容
            style: 表格风格
        """
        soup = BeautifulSoup(html_content, 'html.parser')

        # 只处理body下的直接子元素，避免处理表格内部的元素
        body = soup.find('body')
        if not body:
            # 如果没有body标签，就处理整个文档
            body = soup

        # 统计HTML中的图片数量
        all_imgs = soup.find_all('img')
        print(f"[Word导出] HTML中总共有 {len(all_imgs)} 张图片")
        for idx, img in enumerate(all_imgs):
            img_src = img.get('src', '')
            print(f"[Word导出] 图片{idx}: {img_src[:100] if img_src else 'NO SRC'}...")
            # 检查图片是否在表格内
            parent = img.parent
            in_table = False
            while parent:
                if parent.name == 'table':
                    in_table = True
                    break
                parent = parent.parent
            print(f"[Word导出] 图片{idx}在表格内: {in_table}")

        # 递归处理所有顶层元素
        self._process_elements(doc, body.children, inside_table=False, style=style)

    def _process_elements(self, doc: Document, elements, inside_table=False, style: str = "blue_white"):
        """
        递归处理HTML元素

        Args:
            doc: Word文档对象
            elements: HTML元素列表
            inside_table: 是否在表格内部
            style: 表格风格
        """
        for element in elements:
            # 跳过文本节点和注释
            if not hasattr(element, 'name'):
                continue

            if element.name == 'h1':
                # 一级标题
                if not inside_table:
                    heading = doc.add_heading(element.get_text(), level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif element.name == 'h2':
                # 二级标题
                if not inside_table:
                    heading = doc.add_heading(element.get_text(), level=2)

            elif element.name == 'h3':
                # 三级标题
                if not inside_table:
                    doc.add_heading(element.get_text(), level=3)

            elif element.name == 'p':
                # 段落 - 检查是否在表格内
                if self._is_inside_table(element):
                    continue  # 跳过表格内的段落

                text = element.get_text().strip()
                if text:  # 只添加非空段落
                    para = doc.add_paragraph(text)
                    # 检查对齐方式
                    element_style = element.get('style', '')
                    if 'text-align: center' in element_style or 'text-align:center' in element_style:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif 'text-align: right' in element_style or 'text-align:right' in element_style:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            elif element.name == 'table':
                # 表格 - 检查是否是嵌套表格
                if self._is_nested_table(element):
                    # 嵌套表格：展开处理
                    self._process_nested_table(doc, element, style=style)
                else:
                    # 普通表格
                    self._add_table_to_word(doc, element, style=style)

            elif element.name == 'hr':
                # 分隔线
                if not inside_table:
                    doc.add_paragraph('_' * 50)

            elif element.name == 'img':
                # 图片
                print(f"[Word导出] 发现img标签，inside_table={inside_table}")
                if not inside_table:
                    print(f"[Word导出] 处理img标签")
                    self._add_image_to_word(doc, element)
                else:
                    print(f"[Word导出] 跳过表格内的img标签（应由表格处理器处理）")

            elif element.name in ['div', 'section', 'article']:
                # 容器元素，递归处理其子元素
                print(f"[Word导出] 处理{element.name}容器，子元素数: {len(list(element.children))}")
                self._process_elements(doc, element.children, inside_table, style=style)

    def _add_image_to_word(self, doc: Document, img_element):
        """
        将HTML图片添加到Word文档

        Args:
            doc: Word文档对象
            img_element: BeautifulSoup图片元素
        """
        import requests
        import io
        from urllib.parse import urlparse

        try:
            # 获取图片URL
            img_src = img_element.get('src', '')
            print(f"[Word导出] img_element.get('src'): {img_src[:100] if img_src else 'EMPTY'}")

            if not img_src:
                print(f"[Word导出] 图片没有src属性，跳过")
                return

            print(f"[Word导出] 处理图片: {img_src[:100]}...")

            # 检查是否是base64图片
            if img_src.startswith('data:image'):
                # Base64图片
                import base64
                import re

                # 提取base64数据 - 使用更灵活的方式处理
                # 格式: data:image/png;base64,<base64数据>
                try:
                    # 找到逗号位置，分离header和base64数据
                    comma_index = img_src.find(',')
                    if comma_index == -1:
                        print(f"[Word导出] base64图片格式错误，找不到逗号分隔符")
                        return

                    header = img_src[:comma_index]  # data:image/png;base64
                    base64_data = img_src[comma_index + 1:]  # base64数据

                    # 从header中提取图片格式
                    format_match = re.search(r'image/(\w+)', header)
                    image_format = format_match.group(1) if format_match else 'png'

                    print(f"[Word导出] base64图片格式: {image_format}, 数据长度: {len(base64_data)}")

                    if not base64_data or len(base64_data) == 0:
                        print(f"[Word导出] base64数据为空")
                        return

                    # 解码base64数据
                    image_data = base64.b64decode(base64_data)
                    image_stream = io.BytesIO(image_data)

                    # 添加图片到Word
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(image_stream, width=Inches(5))
                    print(f"[Word导出] 成功添加base64图片，大小: {len(image_data)} 字节")

                except Exception as base64_error:
                    print(f"[Word导出] base64解码失败: {str(base64_error)}")
                    raise

            elif img_src.startswith('http://') or img_src.startswith('https://'):
                # 网络图片
                response = requests.get(img_src, timeout=10)
                if response.status_code == 200:
                    image_stream = io.BytesIO(response.content)

                    # 添加图片到Word
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(image_stream, width=Inches(5))
                    print(f"[Word导出] 成功添加网络图片")
                else:
                    print(f"[Word导出] 下载图片失败，状态码: {response.status_code}")

            else:
                # 本地文件路径（相对或绝对）
                # 注意：这种情况下图片可能无法访问，因为路径可能是前端的路径
                print(f"[Word导出] 跳过本地路径图片: {img_src}")

        except Exception as e:
            print(f"[Word导出] 添加图片到Word失败: {str(e)}")
            import traceback
            traceback.print_exc()

    def _add_image_to_word_cell(self, cell, img_element):
        """
        将HTML图片添加到Word表格单元格

        Args:
            cell: python-docx表格单元格对象
            img_element: BeautifulSoup图片元素
        """
        import requests
        import io
        import base64
        import re
        from urllib.parse import urlparse

        try:
            # 获取图片URL
            img_src = img_element.get('src', '')
            print(f"[Word导出] 表格单元格图片src: {img_src[:100] if img_src else 'EMPTY'}")

            if not img_src:
                print(f"[Word导出] 表格单元格图片没有src属性，跳过")
                return

            # 检查是否是base64图片
            if img_src.startswith('data:image'):
                # Base64图片
                try:
                    # 找到逗号位置，分离header和base64数据
                    comma_index = img_src.find(',')
                    if comma_index == -1:
                        print(f"[Word导出] 表格单元格base64图片格式错误，找不到逗号分隔符")
                        return

                    header = img_src[:comma_index]
                    base64_data = img_src[comma_index + 1:]

                    # 从header中提取图片格式
                    format_match = re.search(r'image/(\w+)', header)
                    image_format = format_match.group(1) if format_match else 'png'

                    print(f"[Word导出] 表格单元格base64图片格式: {image_format}, 数据长度: {len(base64_data)}")

                    if not base64_data or len(base64_data) == 0:
                        print(f"[Word导出] 表格单元格base64数据为空")
                        return

                    # 解码base64数据
                    image_data = base64.b64decode(base64_data)
                    image_stream = io.BytesIO(image_data)

                    # 添加图片到单元格
                    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(image_stream, width=Inches(3))  # 单元格内图片宽度较小
                    print(f"[Word导出] 成功添加表格单元格base64图片，大小: {len(image_data)} 字节")

                except Exception as base64_error:
                    print(f"[Word导出] 表格单元格base64解码失败: {str(base64_error)}")
                    raise

            elif img_src.startswith('http://') or img_src.startswith('https://'):
                # 网络图片
                response = requests.get(img_src, timeout=10)
                if response.status_code == 200:
                    image_stream = io.BytesIO(response.content)

                    # 添加图片到单元格
                    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(image_stream, width=Inches(3))
                    print(f"[Word导出] 成功添加表格单元格网络图片")
                else:
                    print(f"[Word导出] 表格单元格下载图片失败，状态码: {response.status_code}")

            else:
                # 本地文件路径
                print(f"[Word导出] 跳过表格单元格本地路径图片: {img_src}")

        except Exception as e:
            print(f"[Word导出] 添加图片到表格单元格失败: {str(e)}")
            import traceback
            traceback.print_exc()

    def _is_inside_table(self, element) -> bool:
        """
        检查元素是否在表格内部

        Args:
            element: BeautifulSoup元素

        Returns:
            是否在表格内
        """
        parent = element.parent
        while parent:
            if parent.name == 'table':
                return True
            parent = parent.parent
        return False

    def _is_nested_table(self, table_element) -> bool:
        """
        检查表格是否包含嵌套的表格（表格单元格中包含其他表格）

        Args:
            table_element: BeautifulSoup表格元素

        Returns:
            是否是嵌套表格
        """
        # 查找表格内的所有td/th单元格
        cells = table_element.find_all(['td', 'th'])
        for cell in cells:
            # 检查单元格内是否包含表格
            if cell.find('table'):
                return True
        return False

    def _get_cell_background_color(self, style: str, is_label: bool, row_idx: int) -> str:
        """
        根据风格获取单元格背景色

        Args:
            style: 表格风格
            is_label: 是否是标签列
            row_idx: 行索引

        Returns:
            背景色（十六进制，如 'D9E2F3'），如果不需要背景色则返回 None
        """
        if style == "plain":
            # 纯白风格：无背景色
            return None
        elif style == "classic":
            # 经典风格：标签列深蓝色，值列白色
            if is_label:
                return '4472C4'  # 深蓝色
            else:
                return None
        else:  # blue_white（默认）
            # 蓝白相间风格：标签列浅蓝色，值列交替
            if is_label:
                return 'D9E2F3'  # 浅蓝色
            else:
                if row_idx % 2 == 1:  # 奇数行使用浅蓝色
                    return 'EBF1FA'  # 更浅的蓝色
                else:
                    return None

    def _process_nested_table(self, doc: Document, table_element, style: str = "blue_white"):
        """
        处理嵌套表格：将并列的模块转换为一个大表格，每个模块占一列

        Args:
            doc: Word文档对象
            table_element: BeautifulSoup表格元素
        """
        print(f"[Word导出] 检测到嵌套表格，转换为并列表格")

        # 获取外层表格的第一行（通常并列模块只有一行）
        tbody = table_element.find('tbody')
        if tbody:
            outer_rows = tbody.find_all('tr', recursive=False)
        else:
            outer_rows = table_element.find_all('tr', recursive=False)

        if not outer_rows:
            print(f"[Word导出] 嵌套表格没有行，跳过")
            return

        # 只处理第一行（并列模块通常在一行中）
        first_row = outer_rows[0]
        outer_cells = first_row.find_all(['td', 'th'], recursive=False)
        num_modules = len(outer_cells)

        if num_modules == 0:
            print(f"[Word导出] 嵌套表格没有列，跳过")
            return

        print(f"[Word导出] 嵌套表格有 {num_modules} 个并列模块")

        # 收集每个模块的数据
        modules_data = []
        max_rows = 0

        for col_idx, cell in enumerate(outer_cells):
            module_title = None
            module_rows = []

            # 查找单元格内的所有元素
            for element in cell.children:
                if not hasattr(element, 'name'):
                    continue

                if element.name == 'h3':
                    module_title = element.get_text().strip()
                    print(f"[Word导出]   模块 {col_idx + 1} 标题: {module_title}")

                elif element.name == 'div':
                    inner_table = element.find('table')
                    if inner_table:
                        rows = inner_table.find_all('tr')
                        for row in rows:
                            cells = row.find_all(['td', 'th'])
                            if len(cells) >= 2:
                                label = cells[0].get_text().strip()
                                value = cells[1].get_text().strip()
                                module_rows.append((label, value))

                elif element.name == 'table':
                    rows = element.find_all('tr')
                    for row in rows:
                        cells = row.find_all(['td', 'th'])
                        if len(cells) >= 2:
                            label = cells[0].get_text().strip()
                            value = cells[1].get_text().strip()
                            module_rows.append((label, value))

            modules_data.append({
                'title': module_title or f'模块 {col_idx + 1}',
                'rows': module_rows
            })
            max_rows = max(max_rows, len(module_rows))
            print(f"[Word导出]   模块 {col_idx + 1} 有 {len(module_rows)} 行数据")

        # 创建Word表格：标题行 + 数据行
        # 每个模块占2列（标签列 + 值列）
        total_cols = num_modules * 2
        total_rows = max_rows + 1  # +1 for title row

        word_table = doc.add_table(rows=total_rows, cols=total_cols)
        word_table.style = 'Table Grid'

        # 填充表格
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        for module_idx, module_data in enumerate(modules_data):
            col_offset = module_idx * 2

            # 标题行：合并两列
            title_cell = word_table.rows[0].cells[col_offset]
            if col_offset + 1 < total_cols:
                title_cell.merge(word_table.rows[0].cells[col_offset + 1])

            # 设置标题
            title_cell.text = module_data['title']
            for paragraph in title_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                paragraph.alignment = 1  # 居中对齐

            # 数据行
            for row_idx, (label, value) in enumerate(module_data['rows']):
                actual_row_idx = row_idx + 1

                # 标签列
                label_cell = word_table.rows[actual_row_idx].cells[col_offset]
                label_cell.text = label
                for paragraph in label_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

                # 添加标签列的背景色
                label_bg_color = self._get_cell_background_color(style, is_label=True, row_idx=row_idx)
                if label_bg_color:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), label_bg_color)
                    label_cell._element.get_or_add_tcPr().append(shading_elm)

                # 值列
                value_cell = word_table.rows[actual_row_idx].cells[col_offset + 1]
                value_cell.text = value
                for paragraph in value_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

                # 添加值列的背景色
                value_bg_color = self._get_cell_background_color(style, is_label=False, row_idx=row_idx)
                if value_bg_color:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), value_bg_color)
                    value_cell._element.get_or_add_tcPr().append(shading_elm)

        print(f"[Word导出] 创建了 {total_rows}行 x {total_cols}列 的并列表格")
    


    def _add_table_to_word(self, doc: Document, table_element, style: str = "blue_white"):
        """
        将HTML表格添加到Word文档，支持colspan和rowspan

        Args:
            doc: Word文档对象
            table_element: BeautifulSoup表格元素
            style: 表格风格
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        rows = table_element.find_all('tr')
        if not rows:
            return

        # 计算实际的列数（考虑colspan）
        max_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'])
            col_count = 0
            for cell in cells:
                colspan = int(cell.get('colspan', 1))
                col_count += colspan
            max_cols = max(max_cols, col_count)

        if max_cols == 0:
            return

        # 创建表格
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'

        # 填充数据并处理合并单元格
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            col_idx = 0

            for cell in cells:
                if row_idx >= len(table.rows):
                    break

                # 获取colspan和rowspan
                colspan = int(cell.get('colspan', 1))
                rowspan = int(cell.get('rowspan', 1))

                # 确保不超出表格范围
                if col_idx >= len(table.rows[row_idx].cells):
                    break

                # 设置单元格文本和内容
                cell_text = cell.get_text().strip()
                word_cell = table.rows[row_idx].cells[col_idx]

                # 清空单元格内容
                word_cell.text = ''

                # 处理单元格内的图片
                cell_imgs = cell.find_all('img')
                if cell_imgs:
                    print(f"[Word导出] 表格单元格[{row_idx},{col_idx}]中发现 {len(cell_imgs)} 张图片")
                    for img in cell_imgs:
                        self._add_image_to_word_cell(word_cell, img)

                # 添加文本
                if cell_text:
                    para = word_cell.paragraphs[0] if word_cell.paragraphs else word_cell.add_paragraph()
                    para.text = cell_text

                # 表头加粗
                if cell.name == 'th':
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # 设置字体大小
                for paragraph in word_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

                # 设置背景色（假设第一列是标签列）
                is_label = (col_idx == 0)
                bg_color = self._get_cell_background_color(style, is_label=is_label, row_idx=row_idx)
                if bg_color:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), bg_color)
                    word_cell._element.get_or_add_tcPr().append(shading_elm)

                # 第一列加粗
                if col_idx == 0:
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # 处理colspan（横向合并）
                if colspan > 1:
                    try:
                        # 合并当前单元格和右侧的单元格
                        end_col = min(col_idx + colspan - 1, len(table.rows[row_idx].cells) - 1)
                        if end_col > col_idx:
                            word_cell.merge(table.rows[row_idx].cells[end_col])
                    except Exception as e:
                        print(f"[Word导出] 合并列失败: {e}")

                # 处理rowspan（纵向合并）
                if rowspan > 1:
                    try:
                        # 合并当前单元格和下方的单元格
                        end_row = min(row_idx + rowspan - 1, len(table.rows) - 1)
                        if end_row > row_idx:
                            word_cell.merge(table.rows[end_row].cells[col_idx])
                    except Exception as e:
                        print(f"[Word导出] 合并行失败: {e}")

                # 移动到下一列
                col_idx += colspan
    
    def _generate_pdf_html(self, wps: WPS, content: str) -> str:
        """
        生成用于PDF导出的完整HTML
        
        Args:
            wps: WPS对象
            content: HTML内容
            
        Returns:
            完整的HTML字符串
        """
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                    @bottom-center {{
                        content: "第 " counter(page) " 页";
                    }}
                }}
                body {{
                    font-family: 'SimSun', 'Microsoft YaHei', 'Arial', sans-serif;
                    font-size: 12pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    text-align: center;
                    font-size: 24pt;
                    margin-bottom: 10pt;
                    color: #000;
                }}
                h2 {{
                    font-size: 18pt;
                    background-color: #f0f0f0;
                    padding: 8pt;
                    margin-top: 15pt;
                    color: #000;
                }}
                h3 {{
                    font-size: 14pt;
                    margin-top: 10pt;
                    color: #000;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 10pt 0;
                    page-break-inside: avoid;
                }}
                td, th {{
                    border: 1px solid #000;
                    padding: 6pt 8pt;
                    text-align: left;
                }}
                th {{
                    background-color: #f0f0f0;
                    font-weight: bold;
                }}
                img {{
                    max-width: 100%;
                    page-break-inside: avoid;
                }}
                p {{
                    margin: 0.5em 0;
                }}
                hr {{
                    border: none;
                    border-top: 2px solid #ddd;
                    margin: 2em 0;
                }}
            </style>
        </head>
        <body>
            {content}
        </body>
        </html>
        """
    
    def _generate_default_html(self, wps: WPS) -> str:
        """
        生成默认的HTML内容（当document_html为空时）
        
        Args:
            wps: WPS对象
            
        Returns:
            HTML字符串
        """
        return f"""
        <h1>{wps.title or 'WPS文档'}</h1>
        <p style="text-align: center;">文档编号: {wps.wps_number} | 版本: {wps.revision or 'A'}</p>
        <hr />
        <h2>基本信息</h2>
        <table>
            <tbody>
                <tr>
                    <td style="width: 30%; font-weight: bold;">WPS编号</td>
                    <td style="width: 70%;">{wps.wps_number}</td>
                </tr>
                <tr>
                    <td style="width: 30%; font-weight: bold;">标题</td>
                    <td style="width: 70%;">{wps.title or '-'}</td>
                </tr>
                <tr>
                    <td style="width: 30%; font-weight: bold;">版本</td>
                    <td style="width: 70%;">{wps.revision or 'A'}</td>
                </tr>
                <tr>
                    <td style="width: 30%; font-weight: bold;">状态</td>
                    <td style="width: 70%;">{wps.status or '-'}</td>
                </tr>
            </tbody>
        </table>
        <p></p>
        <p style="text-align: center; color: #999; margin-top: 2em;">
            此文档尚未使用文档编辑器编辑，显示的是默认内容。
        </p>
        """

    # ==================== PQR导出方法 ====================

    def export_pqr_to_word(self, pqr: PQR, style: str = "blue_white") -> io.BytesIO:
        """
        导出PQR为Word文档

        Args:
            pqr: PQR对象
            style: 表格风格，可选值：
                - "blue_white": 蓝白相间风格（默认）
                - "plain": 纯白风格
                - "classic": 经典风格（深蓝标题）

        Returns:
            Word文档的BytesIO流
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx未安装，请运行: pip install python-docx")

        # 创建Word文档
        doc = Document()

        # 设置页面（A4）
        section = doc.sections[0]
        section.page_height = Inches(11.69)  # A4高度
        section.page_width = Inches(8.27)    # A4宽度
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)

        # 获取HTML内容
        html_content = getattr(pqr, 'document_html', None) or self._generate_default_pqr_html(pqr)

        # 解析HTML并转换为Word
        self._html_to_word(doc, html_content, style=style)

        # 添加页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"打印日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return file_stream

    def export_pqr_to_pdf(self, pqr: PQR) -> io.BytesIO:
        """
        导出PQR为PDF文档

        Args:
            pqr: PQR对象

        Returns:
            PDF文档的BytesIO流
        """
        if not WEASYPRINT_AVAILABLE:
            raise ImportError("weasyprint未安装，请运行: pip install weasyprint")

        # 获取HTML内容
        html_content = getattr(pqr, 'document_html', None) or self._generate_default_pqr_html(pqr)

        # 生成完整的HTML文档
        full_html = self._generate_pdf_html_for_pqr(pqr, html_content)

        # 生成PDF
        pdf_bytes = HTML(string=full_html).write_pdf()

        return io.BytesIO(pdf_bytes)

    def _generate_default_pqr_html(self, pqr: PQR) -> str:
        """
        生成默认的PQR HTML内容（当document_html为空时）

        Args:
            pqr: PQR对象

        Returns:
            HTML字符串
        """
        test_date = pqr.test_date.strftime('%Y-%m-%d') if pqr.test_date else '-'

        return f"""
        <h1>{pqr.title or 'PQR文档'}</h1>
        <p style="text-align: center;">文档编号: {pqr.pqr_number}</p>
        <hr />
        <h2>基本信息</h2>
        <table>
            <tbody>
                <tr>
                    <td style="width: 30%; font-weight: bold;">PQR编号</td>
                    <td style="width: 70%;">{pqr.pqr_number}</td>
                </tr>
                <tr>
                    <td style="width: 30%; font-weight: bold;">标题</td>
                    <td style="width: 70%;">{pqr.title or '-'}</td>
                </tr>
                <tr>
                    <td style="width: 30%; font-weight: bold;">试验日期</td>
                    <td style="width: 70%;">{test_date}</td>
                </tr>
                <tr>
                    <td style="width: 30%; font-weight: bold;">状态</td>
                    <td style="width: 70%;">{pqr.status or '-'}</td>
                </tr>
                <tr>
                    <td style="width: 30%; font-weight: bold;">评定结果</td>
                    <td style="width: 70%;">{pqr.qualification_result or '-'}</td>
                </tr>
            </tbody>
        </table>
        <p></p>
        <p style="text-align: center; color: #999; margin-top: 2em;">
            此文档尚未使用文档编辑器编辑，显示的是默认内容。
        </p>
        """

    def _generate_pdf_html_for_pqr(self, pqr: PQR, content: str) -> str:
        """
        生成用于PQR PDF导出的完整HTML

        Args:
            pqr: PQR对象
            content: HTML内容

        Returns:
            完整的HTML字符串
        """
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                    @bottom-center {{
                        content: "第 " counter(page) " 页";
                    }}
                }}
                body {{
                    font-family: 'SimSun', 'Microsoft YaHei', 'Arial', sans-serif;
                    font-size: 12pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    text-align: center;
                    font-size: 24pt;
                    margin-bottom: 10pt;
                    color: #000;
                }}
                h2 {{
                    font-size: 18pt;
                    background-color: #f0f0f0;
                    padding: 8pt;
                    margin-top: 15pt;
                    color: #000;
                }}
                h3 {{
                    font-size: 14pt;
                    margin-top: 10pt;
                    color: #000;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 10pt 0;
                    page-break-inside: avoid;
                }}
                td, th {{
                    border: 1px solid #000;
                    padding: 6pt 8pt;
                    text-align: left;
                }}
                th {{
                    background-color: #f0f0f0;
                    font-weight: bold;
                }}
                img {{
                    max-width: 100%;
                    page-break-inside: avoid;
                }}
                p {{
                    margin: 0.5em 0;
                }}
                hr {{
                    border: none;
                    border-top: 2px solid #ddd;
                    margin: 2em 0;
                }}
            </style>
        </head>
        <body>
            {content}
        </body>
        </html>
        """

    # ==================== pPQR导出方法 ====================

    def export_ppqr_to_word(self, ppqr: PPQR, style: str = "blue_white") -> io.BytesIO:
        """
        导出pPQR为Word文档

        Args:
            ppqr: PPQR对象
            style: 表格风格，可选值：
                - "blue_white": 蓝白相间风格（默认）
                - "plain": 纯白风格
                - "classic": 经典风格（深蓝标题）

        Returns:
            Word文档的BytesIO流
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx未安装，请运行: pip install python-docx")

        try:
            # 创建Word文档
            doc = Document()

            # 设置页面（A4）
            section = doc.sections[0]
            section.page_height = Inches(11.69)  # A4高度
            section.page_width = Inches(8.27)    # A4宽度
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
            section.left_margin = Inches(0.79)
            section.right_margin = Inches(0.79)

            # 获取HTML内容
            html_content = getattr(ppqr, 'document_html', None) or self._generate_default_ppqr_html(ppqr)

            # 解析HTML并转换为Word
            self._html_to_word(doc, html_content, style=style)

            # 添加页脚
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = f"打印日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 保存到内存
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return file_stream
        except Exception as e:
            print(f"[pPQR导出Word错误] pPQR编号: {ppqr.ppqr_number}")
            print(f"[pPQR导出Word错误] 错误信息: {str(e)}")
            print(f"[pPQR导出Word错误] 错误类型: {type(e).__name__}")
            import traceback
            print(f"[pPQR导出Word错误] 堆栈跟踪:\n{traceback.format_exc()}")
            raise

    def export_ppqr_to_pdf(self, ppqr: PPQR) -> io.BytesIO:
        """
        导出pPQR为PDF文档

        Args:
            ppqr: PPQR对象

        Returns:
            PDF文档的BytesIO流
        """
        if not WEASYPRINT_AVAILABLE:
            raise ImportError("weasyprint未安装，请运行: pip install weasyprint")

        # 获取HTML内容
        html_content = getattr(ppqr, 'document_html', None) or self._generate_default_ppqr_html(ppqr)

        # 生成完整的HTML文档
        full_html = self._generate_pdf_html_for_ppqr(ppqr, html_content)

        # 生成PDF
        pdf_bytes = HTML(string=full_html).write_pdf()

        return io.BytesIO(pdf_bytes)

    def _generate_default_ppqr_html(self, ppqr: PPQR) -> str:
        """
        生成默认的pPQR HTML内容（当document_html为空时）

        Args:
            ppqr: PPQR对象

        Returns:
            HTML字符串
        """
        try:
            # 安全地获取日期字段
            planned_date = '-'
            actual_date = '-'

            if hasattr(ppqr, 'planned_test_date') and ppqr.planned_test_date:
                planned_date = ppqr.planned_test_date.strftime('%Y-%m-%d')

            if hasattr(ppqr, 'actual_test_date') and ppqr.actual_test_date:
                actual_date = ppqr.actual_test_date.strftime('%Y-%m-%d')

            return f"""
            <h1>{ppqr.title or 'pPQR文档'}</h1>
            <p style="text-align: center;">文档编号: {ppqr.ppqr_number} | 版本: {ppqr.revision or 'A'}</p>
            <hr />
            <h2>基本信息</h2>
            <table>
                <tbody>
                    <tr>
                        <td style="width: 30%; font-weight: bold;">pPQR编号</td>
                        <td style="width: 70%;">{ppqr.ppqr_number}</td>
                    </tr>
                    <tr>
                        <td style="width: 30%; font-weight: bold;">标题</td>
                        <td style="width: 70%;">{ppqr.title or '-'}</td>
                    </tr>
                    <tr>
                        <td style="width: 30%; font-weight: bold;">版本</td>
                        <td style="width: 70%;">{ppqr.revision or 'A'}</td>
                    </tr>
                    <tr>
                        <td style="width: 30%; font-weight: bold;">计划测试日期</td>
                        <td style="width: 70%;">{planned_date}</td>
                    </tr>
                    <tr>
                        <td style="width: 30%; font-weight: bold;">实际测试日期</td>
                        <td style="width: 70%;">{actual_date}</td>
                    </tr>
                    <tr>
                        <td style="width: 30%; font-weight: bold;">状态</td>
                        <td style="width: 70%;">{ppqr.status or '-'}</td>
                    </tr>
                </tbody>
            </table>
            <p></p>
            <p style="text-align: center; color: #999; margin-top: 2em;">
                此文档尚未使用文档编辑器编辑，显示的是默认内容。
            </p>
            """
        except Exception as e:
            print(f"[生成默认pPQR HTML错误] {str(e)}")
            # 返回最简单的HTML
            return f"""
            <h1>{getattr(ppqr, 'title', 'pPQR文档')}</h1>
            <p>文档编号: {getattr(ppqr, 'ppqr_number', 'N/A')}</p>
            """

    def _generate_pdf_html_for_ppqr(self, ppqr: PPQR, content: str) -> str:
        """
        生成用于pPQR PDF导出的完整HTML

        Args:
            ppqr: PPQR对象
            content: HTML内容

        Returns:
            完整的HTML字符串
        """
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                    @bottom-center {{
                        content: "第 " counter(page) " 页";
                    }}
                }}
                body {{
                    font-family: 'SimSun', 'Microsoft YaHei', 'Arial', sans-serif;
                    font-size: 12pt;
                    line-height: 1.6;
                    color: #333;
                }}
                h1 {{
                    text-align: center;
                    font-size: 24pt;
                    margin-bottom: 10pt;
                    color: #000;
                }}
                h2 {{
                    font-size: 18pt;
                    background-color: #f0f0f0;
                    padding: 8pt;
                    margin-top: 15pt;
                    color: #000;
                }}
                h3 {{
                    font-size: 14pt;
                    margin-top: 10pt;
                    color: #000;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 10pt 0;
                    page-break-inside: avoid;
                }}
                td, th {{
                    border: 1px solid #000;
                    padding: 6pt 8pt;
                    text-align: left;
                }}
                th {{
                    background-color: #f0f0f0;
                    font-weight: bold;
                }}
                img {{
                    max-width: 100%;
                    page-break-inside: avoid;
                }}
                p {{
                    margin: 0.5em 0;
                }}
                hr {{
                    border: none;
                    border-top: 2px solid #ddd;
                    margin: 2em 0;
                }}
            </style>
        </head>
        <body>
            {content}
        </body>
        </html>
        """

